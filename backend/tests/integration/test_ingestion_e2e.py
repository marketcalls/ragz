import io

import pytest
from docx import Document as DocxBuilder
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from raghub.modules.documents.ingest import run_chunk, run_embed_upsert, run_parse
from raghub.modules.documents.models import IngestJob
from raghub.modules.documents.pipeline import IngestFailure
from raghub.modules.documents.service import create_from_upload
from raghub.modules.retrieval.service import retrieve, update_document_current
from tests.modules.retrieval.test_retrieve import seed_workspace


def fixture_docx() -> bytes:
    d = DocxBuilder()
    d.add_heading("Operations Manual", level=1)
    for i in range(30):
        d.add_paragraph(f"Routine operational paragraph number {i} about daily procedures.")
    d.add_heading("Power Requirements", level=1)
    d.add_paragraph("The flux capacitor requires exactly 1.21 gigawatts of power "
                    "supplied by the plutonium reactor.")
    d.add_heading("Billing", level=1)
    d.add_paragraph("Invoice 0231 was issued for the October plutonium delivery.")
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


async def test_docx_through_real_pipeline_then_hybrid_retrieval(
    session: AsyncSession, qdrant_collection: None
) -> None:
    ctx, ws = await seed_workspace(session, "e2e1")
    doc = await create_from_upload(session, ctx, ws.id, filename="manual.docx",
                                   mime="application/vnd.openxmlformats-officedocument"
                                        ".wordprocessingml.document",
                                   data=fixture_docx())
    await run_parse(doc.id)   # real Docling (DOCX: no model downloads)
    await run_chunk(doc.id)
    await run_embed_upsert(doc.id)
    await session.refresh(doc)
    assert doc.status == "indexed" and (doc.page_count or 0) >= 1
    stages = {j.stage for j in (await session.execute(
        select(IngestJob).where(IngestJob.document_id == doc.id))).scalars()}
    assert stages == {"parse", "chunk", "embed", "upsert"}

    # Plan H: freshly-upserted points are is_current=False (invisible) until
    # promotion (Task 6). Real promotion doesn't exist yet, so this test
    # stands in for it via the sanctioned update_document_current path.
    await update_document_current(ctx.org_id, doc.id, is_current=True)

    # (a) exact keyword query — real BM25 sparse must surface the billing chunk
    kw = await retrieve(session, ctx, ws.id, "invoice 0231", top_k=5)
    assert any("Invoice 0231" in c.text for c in kw.chunks[:5])

    # (b) partial-overlap query — hybrid fusion puts the power chunk in top-5
    ov = await retrieve(session, ctx, ws.id, "gigawatts flux capacitor power", top_k=5)
    hit = next(c for c in ov.chunks[:5] if "1.21 gigawatts" in c.text)
    assert hit.document_id == doc.id and hit.page >= 1 and hit.chunk_index >= 0


async def test_empty_and_unsupported_fail_cleanly(
    session: AsyncSession, stack_env: None
) -> None:
    ctx, ws = await seed_workspace(session, "e2e2")
    empty = await create_from_upload(session, ctx, ws.id, filename="empty.txt",
                                     mime="text/plain", data=b"")
    with pytest.raises(IngestFailure, match="empty"):
        await run_parse(empty.id)
    await session.refresh(empty)
    assert empty.status == "failed" and "empty" in (empty.error or "")

    weird = await create_from_upload(session, ctx, ws.id, filename="blob.xyz",
                                     mime="application/octet-stream", data=b"\x00\x01")
    with pytest.raises(IngestFailure):
        await run_parse(weird.id)
    await session.refresh(weird)
    assert weird.status == "failed" and weird.error
